from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    doc = Document()

    # --- HALAMAN MUKA ---
    # Sesuai judul dokumen: Akreditasi Program Studi 2.0 Program Magister [cite: 1]
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LAPORAN KINERJA PROGRAM STUDI\n")
    run.bold = True
    run.font.size = Pt(16)
    
    title.add_run("AKREDITASI PROGRAM STUDI 2.0\n")
    title.add_run("PROGRAM MAGISTER\n")
    title.add_run("{{ pt }}\n\n").bold = True
    title.add_run("TAHUN 2025")

    doc.add_page_break()

    # --- IDENTITAS PENGUSUL ---
    # Mengikuti struktur IDENTITAS PENGUSUL di halaman 6 [cite: 58]
    doc.add_heading('IDENTITAS PENGUSUL', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ["Perguruan Tinggi", "{{ pt }}"],
        ["Unit Pengelola Program Studi", "{{ upps }}"],
        ["Jenis Program", "{{ jenis_program }}"],
        ["Nama Program Studi", "{{ nama_prodi }}"],
        ["Alamat", "{{ alamat }}"]
    ]
    
    for i, (label, tag) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = tag

    doc.add_paragraph("\n")

    # --- TABEL PENDANAAN ---
    # Sesuai Tabel 1.A.2 Sumber Pendanaan UPPS/PS [cite: 103]
    doc.add_heading('Tabel 1.A.2 Sumber Pendanaan UPPS/PS', level=2)
    t_dana = doc.add_table(rows=2, cols=4)
    t_dana.style = 'Table Grid'
    
    headers = ['Sumber Pendanaan', 'TS-2', 'TS-1', 'TS']
    for i, h in enumerate(headers):
        t_dana.rows[0].cells[i].text = h
    
    # MENGGUNAKAN TAG FOR STANDAR (Tanpa 'tr')
    # Letakkan {% for %} di sel pertama dan {% endfor %} di sel terakhir pada baris yang sama
    row = t_dana.rows[1].cells
    row[0].text = "{% for d in dana %}{{ d.sumber }}"
    row[1].text = "{{ d.ts2 }}"
    row[2].text = "{{ d.ts1 }}"
    row[3].text = "{{ d.ts }}{% endfor %}"

    doc.save('template_lkps.docx')
    print("Sukses: File 'template_lkps.docx' telah diperbarui dengan tag standar!")

if __name__ == "__main__":
    create_template()